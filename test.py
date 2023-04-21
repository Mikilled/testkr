import re
import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import Inches
from docx.enum.text import WD_LINE_SPACING
import calendar
from datetime import datetime
from telebot import *
import locale
locale.setlocale(
    category=locale.LC_ALL,
    locale="Russian"
)

def iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def pars_info(text):
    fields = [
        (r'[Фф]акультет\s', 'факультет'),
        (r'[Кк]афедра\s', 'кафедра'),
        (r'[Оо]тчет по:', 'тип'),
        (r'[Тт]ема:\s', 'тема'),
        (r'[Дд]исциплина:\s', 'дисциплина'),
        (r'[Вв]ариант:\s', 'вариант'),
        (r'[Гг]руппа:\s', 'группа'),
        (r'[Вв]ыполнил:\s', 'выполнил'),
        (r'[Пп]роверил:\s', 'проверил'),
        (r'[Цц]ель работы:\s', 'цель'),
        (r'[Зз]адание:\s', 'задание')
    ]

    for pattern, key in fields:
        match = re.search(pattern, text)
        if match:
            info[key] = text.split(':', 1)[1].strip()

info = {
    'факультет': "",
    "кафедра": "",
    "тип": "",
    "тема": "",
    "дисциплина": "",
    "вариант": "",
    'группа':"",
    "выполнил": "",
    "проверил": "",
    "цель": "",
    "задание": "",
}




names_tit = []
def pars():
    doc = docx.Document(r"form.docx")
    count = 0
    text = ''
    for i, block in enumerate(iter_block_items(doc)):
        if count == 60:
            break
        count += 1
        if isinstance(block,Paragraph):
            pars_info(block.text)

            text +=block.text

        else:
            for row in block.rows:
                for cell in row.cells:
                    pars_info(cell.text)
                    text += cell.text
    #print(info)



def write_docx(name_file):
    name_vuz = 'МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ\n\nФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ\nОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ\nВЫСШЕГО ОБРАЗОВАНИЯ\n«НОВОСИБИРСКИЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»\n___________________________________________________________________________'
    doc = docx.Document()
    # создаем параграф
    p = doc.add_paragraph()

    section = doc.sections[0]

    # задаем отступы
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    #section.top_margin = Cm(2.54)
    #section.bottom_margin = Cm(2.54)



    # добавляем текст в параграф и устанавливаем выравнивание по центру
    run = p.add_run(name_vuz )
    p.alignment = 1 # for left, 1 for center, 2 right, 3 justify
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    run = p.add_run(f'\n{info["факультет"]}\n{info["кафедра"]}\n\n')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    picture = run.add_picture("nstu_logo.png", width=Inches(3))


    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3

    p.alignment = 1
    run = p.add_run(f'ОТЧЁТ\n')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(16)
    font.bold = True


    run = p.add_run(f'{info["тип"]}\n\n')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True


    run = p.add_run(f'{info["тема"]}\n')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.italic = True


    run = p.add_run('по дисциплине: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    run = p.add_run(f'{info["дисциплина"]}\n')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.italic = True

    if info["вариант"]:
        run = p.add_run(f'Вариант {info["вариант"]}\n')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    else:
        run = p.add_run('\n')

    from docx.table import _Cell
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn


    def set_cell_border(cell: _Cell, **kwargs): #https://www.programmersought.com/article/74085524416/
        """
        Set cell`s border
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))


    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'


    cell1 = table.cell(0, 0)
    set_cell_border(
    cell1,
    top={"sz": 1,"color": "#FFFFFF"},
    bottom={"sz": 1, "color": "#FFFFFF"},
    start={"sz": 1, "color": "#FFFFFF"},
    end={"sz": 1, "color": "#FFFFFF"},
    )
    cell2 = table.cell(0, 1)
    set_cell_border(
    cell2,
    top={"sz": 1,"color": "#FFFFFF"},
    bottom={"sz": 1, "color": "#FFFFFF"},
    start={"sz": 1, "color": "#FFFFFF"},
    end={"sz": 1, "color": "#FFFFFF"},
    )
    cell1_text = cell1.add_paragraph(f'Выполнил(а):\nСтудент гр. {info["группа"]}\n{info["выполнил"]}\n«{datetime.now().strftime("%d")}» {calendar.month_name[datetime.now().month].replace("ь", "я")} {datetime.now().strftime("%Y")}г\n_________________\n(подпись)')
    cell1_text.style.font.name = 'Times New Roman'
    cell1_text.style.font.size = Pt(12)

    cell2_text = cell2.add_paragraph(f'\t\tПроверил:\n\n\t\t{info["проверил"]}\n\t\t«___» ______ {datetime.now().strftime("%Y")}г\n\t\t_________________\n\t\t(подпись)')
    cell2_text.style.font.name = 'Times New Roman'
    cell2_text.style.font.size = Pt(12)



    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(f'\n\n\n\nНовосибирск {datetime.now().strftime("%Y")}г.')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    doc.save(name_file)




pars()
write_docx("testwrite.docx")

# token = '6295000781:AAHo4pljJ1GmkRXRier4-ckD16n8sD5rCjo'
# bot = telebot.TeleBot(token)
# @bot.message_handler(commands=['start'])
# def start(massage):
#     markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
#     upl = types.KeyboardButton("Загрузить отчет")
#     inf = types.KeyboardButton("о боте")
#     markup.add(upl,inf)
#     bot.send_message(massage.chat.id, 'Welcome to the club, buddy', reply_markup=markup)
#
# @bot.message_handler(content_types=['text'])
# def actions(message):
#
#
# bot.polling(none_stop=True)




















