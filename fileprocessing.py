import os

from docx2txt import docx2txt

import docx
from docx.shared import Inches, Mm, RGBColor, Pt
from docx.oxml.table import CT_Tbl
from docx.table import Table
from PIL import Image

def iter_block_items(parent):
    if isinstance(parent, docx.document.Document):  # Исправленное условие
        parent_elm = parent.element.body
    elif isinstance(parent, docx.oxml.table.CT_Tbl):
        parent_elm = parent
    else:
        raise ValueError("Something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
            yield docx.table.Table(child, parent)




def save_text(paragraph,new_paragraph):
    for run in paragraph.runs:
        new_run = new_paragraph.add_run(run.text)

        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.name = 'Times New Roman'
        new_run.font.size = docx.shared.Pt(14)
        new_run.font.color.rgb = RGBColor(0, 0, 0)
    is_list = paragraph.style.name
    print(is_list)




from docx.oxml import OxmlElement

# def get_image_size(doc, rId):
#     rel = doc.part.rels[rId]
#     image_part = rel.target_part
#     width = image_part.image.width
#     height = image_part.image.height
#     width_inch = float(width) / 914400 * Inches(1)
#     height_inch = float(height) / 914400 * Inches(1)
#     print(width_inch, height_inch)
#     return width_inch, height_inch




def process_document(doc_path,new_doc,cout):
    doc = docx.Document(doc_path)
    if not os.path.exists("img_folder/"):
        os.makedirs("img_folder/")
    docx2txt.process(doc_path, 'img_folder/')
    rels = {}
    picture = False
    for r in doc.part.rels.values():
        if isinstance(r._target, docx.parts.image.ImagePart):
            rels[r.rId] = os.path.basename(r._target.partname)
    k = False
    for block in iter_block_items(doc):
        if not k:
            if (not isinstance(block, Table)) and "Новосибирск" in block.text:
                k = True
            continue
        if isinstance(block, docx.text.paragraph.Paragraph):
            if 'Graphic' in block._p.xml:
                for rId in rels:
                    if rId in block._p.xml:
                        image_path = os.path.join('img_folder/', rels[rId])
                        with Image.open(image_path) as img:
                            img_width, img_height = img.size

                        aspect_ratio = img_width / img_height
                        desired_width = 4
                        desired_height = desired_width / aspect_ratio

                        image_paragraph = new_doc.add_paragraph()
                        image_paragraph.add_run().add_picture(image_path, width=Inches(desired_width),height=Inches(desired_height))
                        previous_paragraph = image_paragraph
                        previous_paragraph.paragraph_format.space_after = Pt(0)
                        picture = True
            else:
                if len(block.text.strip()) > 0:
                    for run in block.runs:
                        if all(run.bold for run in block.runs):
                            new_paragraph = new_doc.add_heading()
                            new_paragraph.paragraph_format.space_before = Mm(12)
                            new_paragraph.paragraph_format.space_after = Mm(4)
                            save_text(block, new_paragraph)
                            break
                        elif picture:
                            if ("рис" in block.text.lower()):
                                new_paragraph = new_doc.add_paragraph()
                                paragraph_format = new_paragraph.paragraph_format
                                paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                                run = new_paragraph.add_run(block.text)
                                font = run.font
                                font.italic = True
                            else:
                                print("нет подписи")
                            picture = False
                            if all(run.bold for run in block.runs):
                                new_paragraph = new_doc.add_heading()
                                new_paragraph.paragraph_format.space_before = Mm(12)
                                new_paragraph.paragraph_format.space_after = Mm(4)
                                save_text(block, new_paragraph)
                                break
                            else:
                                new_paragraph = new_doc.add_paragraph()
                                new_paragraph.paragraph_format.line_spacing = 1
                                new_paragraph.paragraph_format.space_before = 0
                                new_paragraph.paragraph_format.space_after = 0
                                save_text(block, new_paragraph)
                                break
                            break
                        else:
                            if block.style.name != "Normal":
                                print(block.text)
                                new_paragraph = new_doc.add_paragraph(style='List Bullet')
                            else:
                                new_paragraph = new_doc.add_paragraph()
                            new_paragraph.paragraph_format.line_spacing = 1
                            new_paragraph.paragraph_format.space_before = 0
                            new_paragraph.paragraph_format.space_after = 0
                            save_text(block, new_paragraph)
                            break


        elif isinstance(block, Table):
            new_table = new_doc.add_table(rows=len(block.rows), cols=len(block.columns))
            new_table.style = 'Table Grid'
            for i, row in enumerate(block.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.cell(i, j)
                    new_paragraph = new_cell.paragraphs[0]
                    new_run = new_paragraph.add_run()
                    new_run.text = cell.text
                    new_run.font.size = Pt(14)
                    new_run.font.name = 'Times New Roman'
                    if i == 0:
                        new_run.font.bold = True

    return  new_doc
    #new_doc.save(out_path)


if __name__ == '__main__':
    new_doc = docx.Document()
    new_doc = process_document(r"docx/gosling228007.docx",new_doc,-1)
    new_doc.save("out-txt.docx")