import os
from docxcompose.composer import Composer
from docx import Document as Document_compose
from parstitul import pars_titul
from writetitul import write_titul
import docx
from docx.shared import Inches, Mm, RGBColor, Pt
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
def iter_block_items(parent):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.oxml.table.CT_Tbl):
        parent_elm = parent
    else:
        raise ValueError("Something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
            yield docx.table.Table(child, parent)



from docx.oxml.shared import OxmlElement, qn
def preventDocumentBreak(document): # https://github.com/python-openxml/python-docx/issues/245
  tags = document.element.xpath('//w:tr')
  rows = len(tags)
  for row in range(0,rows):
    tag = tags[row]
    child = OxmlElement('w:cantSplit')
    tag.append(child)



def save_text(paragraph):
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = docx.shared.Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)


def del_block(block):
    p = block._element
    p.getparent().remove(p)
    p._p = p._element = None



def process_document(doc_path,cout):
    doc = docx.Document(doc_path)
    preventDocumentBreak(doc)
    rels = {}
    count_table = 0
    last_block = 0
    picture = False
    errors = []
    for r in doc.part.rels.values():
        if isinstance(r._target, docx.parts.image.ImagePart):
            rels[r.rId] = os.path.basename(r._target.partname)
    k = False
    for block in iter_block_items(doc):
        if not k:
            if (not isinstance(block, Table)) and "Новосибирск" in block.text:
                k = True
            del_block(block)
            continue


        if isinstance(block, docx.text.paragraph.Paragraph):
            if 'Graphic' in block._p.xml or (isinstance(last_block, Table)):
                for rId in rels:
                    if rId in block._p.xml:
                        # block.paragraph_format.keep_with_next = True
                        block.paragraph_format.space_after = Pt(0)
                        picture = True
            else:
                if len(block.text.strip()) > 0:
                    if '\n' in block.text:
                        block.text = (block.text.replace('\n', ''))
                    for run in block.runs:
                        if all(run.bold for run in block.runs):
                            block.paragraph_format.keep_with_next = True
                            block.paragraph_format.space_before = Mm(12)
                            block.paragraph_format.space_after = Mm(4)
                            save_text(block)
                            break
                        elif picture:
                            if ("рис" in block.text.lower()):
                                # block.text = block.text + '\n'
                                for run1 in block.runs:
                                    run1.font.italic = True
                                paragraph_format = block.paragraph_format
                                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                block.paragraph_format.space_after = Mm(4)
                                last_block.paragraph_format.keep_with_next = True
                                save_text(block)
                            else:
                                if len(block.text.split(' ')) < 7:
                                    block.text = f'Рисунок - ' + block.text
                                    for run1 in block.runs:
                                        run1.font.italic = True
                                    paragraph_format = block.paragraph_format
                                    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    block.paragraph_format.space_after = Mm(4)
                                    last_block.paragraph_format.keep_with_next = True
                                    save_text(block)
                                else:
                                    if "У картинки было обнаружено отсутствие подписи или номер" not in errors:
                                        errors.append("У картинки было обнаружено отсутствие подписи или номер")
                                    #print("нет подписи")
                            picture = False
                            break
                        else:
                            # if block.style.name != "Normal":
                            #     block.paragraph_format.keep_together = True
                            block.paragraph_format.line_spacing = 1.5
                            block.paragraph_format.space_before = 0
                            block.paragraph_format.space_after = 0
                            save_text(block)
                            break
                else:
                    p = block._element
                    p.getparent().remove(p)
                    p._p = p._element = None


        elif isinstance(block, Table):
            count_table+=1
            block.style = 'Table Grid'
            if isinstance(last_block, docx.text.paragraph.Paragraph) and len(last_block.text.split(' ')) < 5:
                last_block.paragraph_format.space_before = Mm(8)
                last_block.paragraph_format.keep_with_next = True
                for run in last_block.runs:
                    run.bold = False
                    run.italic = False
                if "таблица" not in last_block.text.lower():
                    #print("нет номера таблицы")
                    if "У таблицы было обнаружено отсутствие номера" not in errors:
                        errors.append("У таблицы было обнаружено отсутствие номера")
                    last_block.text = f'Таблица {count_table} - ' + last_block.text

            last_row = None
            prev_row = None
            for i, row in enumerate(block.rows):
                for j, cell in enumerate(row.cells):
                    if cell.paragraphs:
                        if cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                            cell.paragraphs[0].runs[0].font.size = Pt(12)
                            if i == 0:
                                cell.paragraphs[0].runs[0].font.bold = True
                                cell.paragraphs[0].paragraph_format.keep_with_next = True
                prev_row = last_row
                last_row = row
            if prev_row is not None:
                for cell in prev_row.cells:
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].paragraph_format.keep_with_next = True
        last_block = block
    return  doc,errors
    #new_doc.save(out_path)

if __name__ == '__main__':
    count, info = pars_titul(r"docx/pr2_final.docx")
    new_doc,err = process_document(r"docx/pr2_final.docx",count)
    new_doc.save(r"docx/2222.docx")
    doc = write_titul("docx/testwrite.docx",info)
    doc.save('docx/testwrite.docx')

    master = Document_compose('docx/testwrite.docx')
    composer = Composer(master)
    doc2 = Document_compose(r"docx/2222.docx")
    composer.append(doc2)
    composer.save("docx/combined.docx")