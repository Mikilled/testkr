import re
import docx
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import Inches
import calendar

import datetime
import locale
from telebot import *
from parstitul import pars_titul
locale.setlocale(
    category=locale.LC_ALL,
    locale="Russian"
)


def write_titul(name_file,info):
    name_vuz = 'МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ\n\nФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ\nОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ\nВЫСШЕГО ОБРАЗОВАНИЯ\n«НОВОСИБИРСКИЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»\n___________________________________________________________________________'
    doc = docx.Document()

    p = doc.add_paragraph()

    section = doc.sections[0]

    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    #section.top_margin = Cm(2.54)
    #section.bottom_margin = Cm(2.54)


    run = p.add_run(name_vuz )
    p.alignment = 1 # for left, 1 for center, 2 right, 3 justify
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    run = p.add_run(f'\n{info["факультет"]}\n{info["кафедра"]}\n\n')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    picture = run.add_picture("nstu_logo.png", width=Inches(3))


    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3

    p.alignment = 1
    run = p.add_run(f'ОТЧЁТ\n')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(16)
    font.bold = True


    run = p.add_run(f'{info["тип"]}\n\n')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True


    run = p.add_run(f'{info["тема"][:2].upper()+info["тема"][2:].lower()}\n')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.italic = True


    run = p.add_run('по дисциплине: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    run = p.add_run(f'{info["дисциплина"]}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.italic = True

    if info["вариант"]:
        run = p.add_run(f'\n{info["вариант"]}\n')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    else:
        run = p.add_run('\n')

    from docx.table import _Cell
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn


    def set_cell_border(cell: _Cell, **kwargs): #https://www.programmersought.com/article/74085524416/

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()


        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)


        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)


                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)


                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))


    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'


    cell1 = table.cell(0, 0)
    cell1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_border(
    cell1,
    top={"sz": 1,"color": "#FFFFFF"},
    bottom={"sz": 1, "color": "#FFFFFF"},
    start={"sz": 1, "color": "#FFFFFF"},
    end={"sz": 1, "color": "#FFFFFF"},
    )
    cell2 = table.cell(0, 1)
    set_cell_border(
    cell2,
    top={"sz": 1,"color": "#FFFFFF"},
    bottom={"sz": 1, "color": "#FFFFFF"},
    start={"sz": 1, "color": "#FFFFFF"},
    end={"sz": 1, "color": "#FFFFFF"},
    )

    cell3 = table.cell(1, 1)
    set_cell_border(
    cell3,
    top={"sz": 1,"color": "#FFFFFF"},
    bottom={"sz": 1, "color": "#FFFFFF"},
    start={"sz": 1, "color": "#FFFFFF"},
    end={"sz": 1, "color": "#FFFFFF"},
    )

    cell4 = table.cell(1, 0)
    set_cell_border(
    cell4,
    top={"sz": 1,"color": "#FFFFFF"},
    bottom={"sz": 1, "color": "#FFFFFF"},
    start={"sz": 1, "color": "#FFFFFF"},
    end={"sz": 1, "color": "#FFFFFF"},
    )

    cell1_text = cell1.add_paragraph(f'Выполнил(а):\nСтудент гр. {info["группа"]}\n{info["выполнил"]}')
    cell1_text.style.font.name = 'Times New Roman'
    cell1_text.style.font.size = Pt(12)
    cell1_text.paragraph_format.space_before = 0
    cell1_text.paragraph_format.space_after = 0



    cell2_text = cell2.add_paragraph(f'\t\tПроверил:\n\n\t\t{info["проверил"]}')
    cell2_text.style.font.name = 'Times New Roman'
    cell2_text.style.font.size = Pt(12)
    cell2_text.paragraph_format.space_before = 0
    cell2_text.paragraph_format.space_after = 0





    cell3_text = cell3.add_paragraph(f'\t\t«___» ______ {datetime.now().strftime("%Y")}г\n\t\t_________________\n\t\t(подпись)')
    cell3_text.style.font.name = 'Times New Roman'
    cell3_text.style.font.size = Pt(13)
    cell3_text.paragraph_format.space_before = 0
    cell3_text.paragraph_format.space_after = 0





    cell4_text = cell4.add_paragraph(f'«{datetime.now().strftime("%d")}» {calendar.month_name[datetime.now().month].replace("ь", "я")} {datetime.now().strftime("%Y")}г\n_________________\n(подпись)')
    cell4_text.style.font.name = 'Times New Roman'
    cell4_text.style.font.size = Pt(13)
    cell4_text.paragraph_format.space_before = 0
    cell4_text.paragraph_format.space_after = 0


    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(f'Новосибирск {datetime.now().strftime("%Y")}г.')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # doc.add_page_break()
    doc.save(name_file)

    print(info)
    return  doc


if __name__ == '__main__':
    count,info =pars_titul(r"C:\Users\admin\Downloads\Otchyot.docx")
    print(count)
    doc = write_titul("testwrite.docx",info)
    doc.save('testwrite.docx')




















