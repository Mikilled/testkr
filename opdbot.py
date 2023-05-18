# from telebot import *
#
#
# token = '5914429354:AAEgxxwOAHRkg0dP-Dq3JII4oGTH5IeHpoo'
# bot = telebot.TeleBot(token)
# @bot.message_handler(commands=['start'])
# def start(massage):
#     markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
#     upl = types.KeyboardButton("Загрузить отчет")
#     inf = types.KeyboardButton("о боте")
#     markup.add(upl,inf)
#     bot.send_message(massage.chat.id, 'Welcome to the club, buddy', reply_markup=markup)
#
# @bot.message_handler(content_types=['text'])
# def actions(message):
#     print(1111)
#
#
# bot.polling(none_stop=True)



from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph): #https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
    # выравниваем параграф по центру
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # запускаем динамическое обновление параграфа
    page_num_run = paragraph.add_run()
    # обозначаем начало позиции вывода
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    # задаем вывод текущего значения страницы PAGE (всего страниц NUMPAGES)
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    # обозначаем конец позиции вывода
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    # добавляем все в наш параграф (который формируется динамически)
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)
#add_page_number(new_doc.sections[0].footer.paragraphs[0])


from docx.table import Table
from PIL import Image
import os
import docx2txt
import docx
from docx.shared import RGBColor
from docx.shared import Mm, Inches
def save_text(paragraph,new_paragraph):
# Копируем текст исходного файла в новый файл
    for run in paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.name = 'Times New Roman'
        new_run.font.size = docx.shared.Pt(14)
        new_run.font.color.rgb = RGBColor(0, 0, 0)





doc = docx.Document('test-txt.docx')
new_doc = docx.Document()
docx2txt.process('test-txt.docx', 'img_folder/')
rels = {}
for r in doc.part.rels.values():
    if isinstance(r._target, docx.parts.image.ImagePart):
        rels[r.rId] = os.path.basename(r._target.partname)
for paragraph in doc.paragraphs:
    if 'Graphic' in paragraph._p.xml:
        for rId in rels:
            if rId in paragraph._p.xml:
                image_path = os.path.join('img_folder', rels[rId])
                with Image.open(image_path) as img:
                    img_width, img_height = img.size

                aspect_ratio = img_width / img_height
                desired_width = 4
                desired_height = desired_width / aspect_ratio
                new_doc.add_picture(image_path, width=Inches(desired_width), height=Inches(desired_height))
                break
    else:
        if len(paragraph.text.strip()) > 0:
            for run in paragraph.runs:
                if all(run.bold for run in paragraph.runs):
                    new_paragraph = new_doc.add_heading()
                    new_paragraph.paragraph_format.space_before = Mm(12)
                    new_paragraph.paragraph_format.space_after = Mm(4)
                    break
                else:
                    new_paragraph = new_doc.add_paragraph()
                    new_paragraph.paragraph_format.line_spacing = 1
                    new_paragraph.paragraph_format.space_before = 0
                    new_paragraph.paragraph_format.space_after = 0
                    break
            save_text(paragraph, new_paragraph)


new_doc.save('out-txt.docx')








# doc = docx.Document('test-txt.docx')
# new_doc = docx.Document()
#
#
# for paragraph in doc.paragraphs:
#     # if len(paragraph.text.strip()) > 0:
#     if True:
#         for run in paragraph.runs:
#             if all(run.bold for run in paragraph.runs):
#                 print(paragraph.text)
#                 new_paragraph = new_doc.add_heading()
#                 new_paragraph.paragraph_format.space_before = Mm(12)
#                 new_paragraph.paragraph_format.space_after = Mm(4)
#                 break
#             else:
#                 new_paragraph = new_doc.add_paragraph()
#                 new_paragraph.paragraph_format.line_spacing = 1
#                 new_paragraph.paragraph_format.space_before = 0
#                 new_paragraph.paragraph_format.space_after = 0
#                 break
#         save_text(paragraph,new_paragraph)
#
# new_doc.save('out-txt.docx')






