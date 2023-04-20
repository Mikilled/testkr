import itertools

import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def getTable(filename):
    wordDoc = docx.Document(filename)
    print(wordDoc.tables)
    for table in wordDoc.tables:
        for row in table.rows:
            for cell in row.cells:
                print (cell.text)



# (getText("test.docx"))

from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def get_data_from_line(num,doc):
    block = iter_block_items(doc)


#функция сбора данных для заполнения титульника
def get_user_data(str):
    if 'факультет' in str.lower():
        print(str)
    elif 'кафедра' in str.lower():
        print(str)
    elif 'практик' in str.lower():
        print(str)
    # elif ('тема' in str.lower() or "«" in str.lower() or '"' in str.lower() or "'" in str.lower()):
    #     print(str)
    elif 'дисциплин' in str.lower():
        print(str)
    elif 'вариант' in str.lower():
        print(str)
    elif 'выполнил' in str.lower():
        print(str)
    elif 'проверил' in str.lower():
        print(str)
    else:
        return



info = {
    'факультет': "",
    "кафедра": "",
    "тип": "",
    "тема": "",
    "дисциплина": "",
    "вариант": "",
    "выполнил": "",
    "проверил": "",
}
import docx
#doc = docx.Document('test.docx')
# print(next(itertools.islice(iter_block_items(doc), 25, None)).text)
# get_user_data(next(itertools.islice(iter_block_items(doc), 16, None)).text)
#
doc = docx.Document('test.docx')
count = 0
for i, block in enumerate(iter_block_items(doc)):
    if count == 30:
        break
    # Ваш код для обработки блока
    count += 1
    if isinstance(block,Paragraph):
        get_user_data(block.text)
       #print(block.text)
        #print("paragraph\n")
    else:
        for row in block.rows:
            for cell in row.cells:
              #print(cell.text)
                #print("table\n")
              get_user_data(cell.text)








def get_para_data(output_doc_name, paragraph):
#https://stackoverflow.com/questions/48869423/how-do-i-copy-the-contents-of-a-word-document
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

# # Imports
# from docx import *
# input_doc = Document('test.docx')
# output_doc = Document()
# # Call the function
# get_para_data(output_doc, input_doc.paragraphs[0])
#
# # Save the new file
# output_doc.save('OutputDoc.docx')
